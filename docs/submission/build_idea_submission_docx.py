from pathlib import Path
from textwrap import wrap
from urllib.request import urlretrieve

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
ASSET_DIR = ROOT / "docs" / "submission" / "generated" / "assets"
OUT_PATH = Path("/storage/emulated/0/Download/Idea-Submission-Mini-Project-Radius-Submission-Ready.docx")
EWU_LOGO_URL = "https://www.ewubd.edu/themes/east-west-university/assets/default/images/logo.png"
EWU_LOGO_PATH = ASSET_DIR / "ewu-logo.png"

FONT = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

GROUP_MEMBERS = [
    ("Md Sadruzzahan Khan", "2023-1-60-147"),
    ("Toyabur Rhaman", "2023-1-60-065"),
    ("Jannat Milky", "2023-1-60-042"),
    ("Rubaiya Zaman Mysha", "2022-1-10-174"),
    ("Al Fahim Fuyad", "2023-1-60-066"),
]


def font(size, bold=False):
    return ImageFont.truetype(BOLD if bold else FONT, size)


def ensure_ewu_logo():
    if not EWU_LOGO_PATH.exists():
        urlretrieve(EWU_LOGO_URL, EWU_LOGO_PATH)
    return EWU_LOGO_PATH


def text_size(draw, text, fnt):
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrapped_text(draw, xy, text, width, fnt, fill=(24, 31, 42), spacing=5, align="center"):
    words = text.split()
    lines = []
    line = ""
    for word in words:
        candidate = f"{line} {word}".strip()
        if text_size(draw, candidate, fnt)[0] <= width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    x, y = xy
    line_height = text_size(draw, "Ag", fnt)[1] + spacing
    for line in lines:
        tw, _ = text_size(draw, line, fnt)
        tx = x + (width - tw) / 2 if align == "center" else x
        draw.text((tx, y), line, font=fnt, fill=fill)
        y += line_height
    return y


def oval(draw, box, text, fill, outline, text_fill=(20, 24, 34)):
    draw.ellipse(box, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    fnt = font(22, True)
    lines = []
    words = text.split()
    line = ""
    max_width = (x2 - x1) - 38
    for word in words:
        candidate = f"{line} {word}".strip()
        if text_size(draw, candidate, fnt)[0] <= max_width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    line_height = 27
    total = len(lines) * line_height
    y = y1 + ((y2 - y1) - total) / 2
    for line in lines:
        tw, _ = text_size(draw, line, fnt)
        draw.text((x1 + ((x2 - x1) - tw) / 2, y), line, font=fnt, fill=text_fill)
        y += line_height


def actor(draw, center, label):
    x, y = center
    stroke = (28, 42, 60)
    draw.ellipse((x - 20, y - 72, x + 20, y - 32), outline=stroke, width=4)
    draw.line((x, y - 32, x, y + 38), fill=stroke, width=4)
    draw.line((x - 45, y - 8, x + 45, y - 8), fill=stroke, width=4)
    draw.line((x, y + 38, x - 42, y + 92), fill=stroke, width=4)
    draw.line((x, y + 38, x + 42, y + 92), fill=stroke, width=4)
    fnt = font(22, True)
    for i, line in enumerate(label.split("\n")):
        tw, _ = text_size(draw, line, fnt)
        draw.text((x - tw / 2, y + 108 + i * 28), line, font=fnt, fill=stroke)


def line(draw, start, end, fill=(66, 84, 108), width=3):
    draw.line((start, end), fill=fill, width=width)


def polyline(draw, points, fill=(66, 84, 108), width=3):
    draw.line(points, fill=fill, width=width, joint="curve")


def dashed(draw, start, end, fill=(112, 130, 154), width=2, dash=14):
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    dist = (dx * dx + dy * dy) ** 0.5
    steps = int(dist / dash)
    for i in range(0, steps, 2):
        a = i / steps
        b = min(i + 1, steps) / steps
        draw.line((x1 + dx * a, y1 + dy * a, x1 + dx * b, y1 + dy * b), fill=fill, width=width)


def label(draw, xy, text, fill=(66, 84, 108)):
    draw.text(xy, text, font=font(18), fill=fill)


def generalization(draw, child, parent):
    line(draw, child, parent, fill=(66, 84, 108), width=3)
    x1, y1 = child
    x2, y2 = parent
    # Hollow triangle pointing toward the parent actor.
    if y2 < y1:
        points = [(x2, y2), (x2 - 13, y2 + 24), (x2 + 13, y2 + 24)]
    else:
        points = [(x2, y2), (x2 - 13, y2 - 24), (x2 + 13, y2 - 24)]
    draw.polygon(points, outline=(66, 84, 108), fill="#f8fafc")


def create_core_use_case(path):
    img = Image.new("RGB", (2200, 1700), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((310, 90, 1890, 1540), radius=24, outline="#334155", width=4, fill="#ffffff")
    draw.text((835, 125), "Radius - Core Marketplace Use Cases", font=font(38, True), fill="#0f172a")

    actor(draw, (130, 330), "Guest")
    actor(draw, (130, 1075), "Registered\nUser")
    actor(draw, (2070, 850), "Admin")
    generalization(draw, (130, 963), (130, 455))

    usecases = {
        "register": (455, 205, 815, 300, "Register / Sign Up"),
        "login": (455, 340, 815, 435, "Login"),
        "browse": (455, 500, 815, 600, "Browse Nearby Listings"),
        "search": (455, 650, 815, 745, "Search & Filter"),
        "detail": (455, 800, 815, 895, "View Listing Details"),
        "radar": (905, 500, 1285, 600, "Trust Radar Discovery"),
        "create": (455, 930, 815, 1030, "Create Listing"),
        "upload": (905, 880, 1285, 980, "Upload Photos"),
        "fraud": (905, 1015, 1285, 1115, "Run Fraud Check"),
        "manage": (455, 1085, 815, 1185, "Edit / Delete Own Listing"),
        "chat": (905, 1195, 1285, 1295, "Proper Chat Inbox"),
        "trade": (1355, 1195, 1715, 1295, "Request / Complete Trade"),
        "report": (455, 1245, 815, 1345, "Report Suspicious Listing"),
        "review": (455, 1395, 815, 1495, "Rate & Review User"),
        "admin_users": (1355, 210, 1715, 305, "Manage User Accounts"),
        "admin_listings": (1355, 350, 1715, 445, "Manage / Remove Listings"),
        "admin_reports": (1355, 490, 1715, 590, "Review Reports & Fraud Queue"),
        "admin_ml": (1355, 1365, 1715, 1460, "Inspect ML Log"),
        "analytics": (1355, 635, 1715, 730, "View Platform Analytics"),
    }
    for key, (x1, y1, x2, y2, text) in usecases.items():
        fill = "#e0f2fe" if key.startswith("admin") or key == "analytics" else "#ecfeff"
        outline = "#2563eb" if key.startswith("admin") or key == "analytics" else "#0f766e"
        if key in {"fraud", "admin_ml"}:
            fill, outline = "#fef3c7", "#d97706"
        oval(draw, (x1, y1, x2, y2), text, fill, outline)

    # Guest associations
    for target in ["register", "login", "browse", "search", "detail"]:
        x1, y1, x2, y2, _ = usecases[target]
        line(draw, (175, 330), (x1, (y1 + y2) / 2))

    # Registered user associations
    registered_routes = {
        "create": [(175, 1045), (455, 980)],
        "manage": [(175, 1085), (455, 1135)],
        "chat": [(175, 1125), (325, 1245), (905, 1245)],
        "report": [(175, 1165), (455, 1295)],
        "review": [(175, 1205), (455, 1445)],
    }
    for points in registered_routes.values():
        polyline(draw, points)

    # Admin associations
    for target in ["admin_users", "admin_listings", "admin_reports", "admin_ml", "analytics"]:
        x1, y1, x2, y2, _ = usecases[target]
        line(draw, (2025, 850), (x2, (y1 + y2) / 2))

    dashed(draw, (815, 550), (905, 550))
    label(draw, (835, 515), "<<include>>")
    dashed(draw, (815, 980), (905, 930))
    label(draw, (825, 900), "<<include>>")
    dashed(draw, (815, 980), (905, 1065))
    label(draw, (825, 1055), "<<include>>")
    dashed(draw, (1285, 1260), (1355, 1260))
    label(draw, (1300, 1175), "<<include>>")
    dashed(draw, (1715, 540), (1715, 1412))
    label(draw, (1728, 965), "view score / signals")

    draw.text((340, 1580), "Notes: Registered User inherits Guest browsing capabilities. Admin decisions create ML labels for future model training.", font=font(24), fill="#475569")
    img.save(path)


def create_ai_use_case(path):
    img = Image.new("RGB", (1800, 1150), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((260, 90, 1540, 1040), radius=24, outline="#334155", width=4, fill="#ffffff")
    draw.text((630, 120), "Radius - AI Fraud Detection Subsystem", font=font(34, True), fill="#0f172a")

    actor(draw, (120, 310), "Registered\nUser")
    actor(draw, (120, 785), "Fraud Detection\nEngine")
    actor(draw, (1680, 640), "Admin")

    usecases = {
        "submit": (410, 190, 760, 285, "Submit Listing for Fraud Check"),
        "score": (785, 305, 1125, 400, "Generate Fraud Risk Score"),
        "image": (380, 500, 700, 595, "Detect Duplicate / Stolen Image"),
        "price": (740, 500, 1060, 595, "Detect Price Anomaly"),
        "behavior": (1100, 500, 1460, 595, "Detect Seller + Text Risk Signals"),
        "snapshot": (740, 675, 1060, 770, "Store Prediction + Feature Snapshot"),
        "queue": (1120, 735, 1465, 830, "View Flagged Listings Queue"),
        "approve": (1120, 875, 1465, 965, "Approve Listing"),
        "remove": (730, 875, 1075, 965, "Remove / Reject Listing"),
        "label": (380, 875, 700, 965, "Create ML Label"),
    }
    for key, (x1, y1, x2, y2, text) in usecases.items():
        fill, outline = "#ecfeff", "#0f766e"
        if key in {"score", "snapshot"}:
            fill, outline = "#fef3c7", "#d97706"
        if key in {"queue", "approve", "remove", "label"}:
            fill, outline = "#ede9fe", "#7c3aed"
        oval(draw, (x1, y1, x2, y2), text, fill, outline)

    line(draw, (165, 310), (410, 238))
    polyline(draw, [(165, 785), (300, 785), (300, 352), (785, 352)])
    line(draw, (1635, 640), (1465, 783))
    line(draw, (1635, 640), (1465, 920))
    line(draw, (1635, 640), (1075, 920))

    dashed(draw, (760, 238), (785, 352))
    label(draw, (720, 280), "<<include>>")
    for target in ["image", "price", "behavior"]:
        x1, y1, x2, y2, _ = usecases[target]
        dashed(draw, (955, 400), ((x1 + x2) / 2, y1))
    label(draw, (945, 445), "<<include>>")
    dashed(draw, (955, 400), (900, 675))
    label(draw, (840, 625), "<<include>>")
    dashed(draw, (1060, 722), (1120, 782))
    label(draw, (1080, 725), "score + signals")
    dashed(draw, (1290, 830), (1290, 875))
    label(draw, (1305, 845), "manual decision")
    dashed(draw, (730, 920), (700, 920))
    label(draw, (690, 880), "label")

    draw.text((285, 1000), "Outputs: score 0-100, threshold band, model version, component scores, explanations, feature snapshot hash.", font=font(22), fill="#475569")
    img.save(path)


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_pair_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell(table.cell(0, 0), rows[0][0], True)
    set_cell(table.cell(0, 1), rows[0][1], True)
    shade(table.cell(0, 0), "D9EAF7")
    shade(table.cell(0, 1), "D9EAF7")
    for left, right in rows[1:]:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    return table


def create_docx(core_image, ai_image):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    logo = ensure_ewu_logo()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo), width=Inches(2.7))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = [
        ("Course Title: Web Programming", 12, False),
        ("Course Code: CSE479", 12, False),
        ("Section: 01", 12, False),
        ("Mini Project Idea Submission", 16, True),
    ]
    for text, size, bold in lines:
        r = p.add_run(text + "\n")
        r.font.size = Pt(size)
        r.bold = bold

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RADIUS\n")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(23, 80, 115)
    r = p.add_run("A Hyperlocal Secondhand Marketplace with AI-Powered Fraud Detection")
    r.bold = True
    r.font.size = Pt(13)

    members = "\n".join(f"{name} - {student_id}" for name, student_id in GROUP_MEMBERS)
    add_pair_table(doc, [
        ("Submitted To", "Group Members"),
        ("Md. Arman Hossain\nLecturer, Department of CSE\nEast West University", members),
    ])
    doc.add_paragraph("Date of Submission: July 05, 2026")
    doc.add_page_break()

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "RADIUS is a hyperlocal secondhand marketplace that helps users buy and sell items near their location. "
        "The project combines nearby product discovery with an explainable AI/ML trust layer. It supports Guest, Registered User, and Admin roles. "
        "Guests can browse nearby listings; registered users can create listings, upload photos, chat, request trades, report suspicious listings, and review users after completed trades; admins can review reports, manage users, inspect fraud queues, and view ML prediction logs."
    )
    doc.add_paragraph(
        "The main intelligent feature is an AI-assisted fraud and fake-listing detection subsystem. When a listing is submitted, the Python/FastAPI trust service analyzes duplicate images, suspicious prices, risky seller behavior, reused text, brand mismatch, prohibited terms, and off-platform payment pressure. It returns a fraud score, decision, explanation list, component scores, model version, and feature snapshot hash for admin transparency."
    )

    doc.add_heading("2. Use Case Diagram 1: Core Marketplace", level=1)
    doc.add_paragraph("This diagram shows the user-facing marketplace features and role-based access in the current Radius implementation.")
    doc.add_picture(str(core_image), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 1: Core Marketplace Use Case Diagram")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("2.1 Actors", level=2)
    add_bullets(doc, [
        "Guest: Can browse, search, filter, and view public listings.",
        "Registered User: Can act as both buyer and seller using the same account.",
        "Admin: Can review reports, manage users, inspect analytics, and moderate suspicious listings.",
    ])

    doc.add_heading("2.2 Core Use Case Descriptions", level=2)
    core_rows = [
        ("Use Case", "Description"),
        ("Register / Sign Up", "Creates a registered user account using name, email, password, and default location."),
        ("Login", "Authenticates users and admins with JWT-based sessions."),
        ("Browse Nearby Listings", "Uses Supabase Postgres with PostGIS to show listings within a configurable radius."),
        ("Trust Radar Discovery", "Displays nearby products as interactive nodes with distance and trust status."),
        ("Create Listing", "Requires product details and at least one server-issued photo upload."),
        ("Run Fraud Check", "Automatically scores the listing with the AI/ML trust service."),
        ("Proper Chat Inbox", "Shows buyer-seller conversations with listing thumbnails and last-message previews."),
        ("Trade Request Flow", "Supports requested, accepted, rejected, cancelled, and completed trade states."),
        ("Admin Moderation", "Allows admins to review reports, inspect risk signals, approve listings, remove listings, and manage users."),
    ]
    add_pair_table(doc, core_rows)

    doc.add_page_break()
    doc.add_heading("3. Use Case Diagram 2: AI Fraud Detection Subsystem", level=1)
    doc.add_paragraph("This diagram shows the intelligent subsystem that analyzes submitted listings and supports admin review.")
    doc.add_picture(str(ai_image), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 2: AI Fraud Detection Subsystem Use Case Diagram")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3.1 AI/ML Use Case Descriptions", level=2)
    ai_rows = [
        ("Use Case", "Description"),
        ("Submit Listing for Fraud Check", "Triggered automatically whenever a registered user creates a listing."),
        ("Generate Fraud Risk Score", "Combines image, price, seller, text, and policy signals into a 0-100 score."),
        ("Detect Duplicate / Stolen Image", "Uses perceptual image hashing and Hamming distance to detect reused product photos."),
        ("Detect Price Anomaly", "Compares price against category-condition market baselines."),
        ("Detect Seller + Text Risk Signals", "Checks new seller risk, prior flags, urgent language, off-platform contact, prohibited words, reused text, and brand mismatch."),
        ("Store Prediction + Feature Snapshot", "Saves model version, component scores, explanations, and feature snapshot hash."),
        ("View Flagged Listings Queue", "Shows suspicious or reported listings to admins with score and signals."),
        ("Approve / Remove Listing", "Admin makes the final decision and creates labels for future ML training."),
    ]
    add_pair_table(doc, ai_rows)

    doc.add_heading("4. Technology Stack", level=1)
    tech_rows = [
        ("Layer", "Technology"),
        ("Frontend", "React, Vite, Framer Motion, Lucide React, custom CSS"),
        ("Backend", "Node.js, Express.js, Socket.io, JWT, bcrypt, Zod"),
        ("Database", "Supabase Postgres with PostGIS geofencing"),
        ("Storage", "Supabase Storage for listing photos"),
        ("AI/ML Service", "Python, FastAPI, scikit-learn, pandas, Pillow, imagehash"),
        ("Testing", "Vitest, Pytest, Vite production build"),
    ]
    add_pair_table(doc, tech_rows)

    doc.add_heading("5. Dataset and Model Plan", level=1)
    doc.add_paragraph(
        "The project includes generated datasets suitable for academic demonstration and model experimentation: 15,000 price rows in market_prices_expanded.csv and 20,000 labeled synthetic fraud rows in fraud_listings_synthetic.csv. The system also includes a collector template for ethical future data collection and a manual CSV import endpoint for labeled examples."
    )
    add_bullets(doc, [
        "Price suggestion uses Random Forest regression.",
        "Candidate fraud model training uses TF-IDF features and Multinomial Naive Bayes.",
        "Evaluation metrics include precision, recall, F1 score, and false positive rate.",
        "Admin decisions create strong labels; user reports create weak labels.",
    ])

    doc.add_heading("6. Alignment With Project Requirements", level=1)
    req_rows = [
        ("Requirement", "Current Radius Implementation"),
        ("Different user roles", "Guest, Registered User, Admin"),
        ("Product selling system", "Listing CRUD, photo upload, manage tab, listing detail, chat, trade flow"),
        ("AI/ML feature", "Fraud scoring, duplicate image detection, price anomaly detection, price suggestion, candidate model training"),
        ("Dataset", "Generated price and fraud datasets plus collector/import pipeline"),
        ("Python learning", "FastAPI ML service using pandas, scikit-learn, Pillow, imagehash"),
        ("Portfolio quality", "Trust Radar, proper inbox, admin dashboard, ML Log, polished UI, tests, docs"),
    ]
    add_pair_table(doc, req_rows)

    doc.add_heading("7. Current Implementation Status", level=1)
    add_bullets(doc, [
        "Live database is populated with showcase users, listings, reports, flagged listings, and trades.",
        "The account istykhan.ik@gmail.com owns showcase listings, including clean and fraud examples.",
        "Admin dashboard is populated with reports, fraud queue items, users, analytics, and ML prediction logs.",
        "Chat works as a proper inbox with conversation previews and selected threads.",
        "The project has automated server tests, ML tests, and client build verification.",
    ])

    doc.save(OUT_PATH)


def main():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    core = ASSET_DIR / "usecase_core_marketplace.png"
    ai = ASSET_DIR / "usecase_ai_fraud.png"
    create_core_use_case(core)
    create_ai_use_case(ai)
    create_docx(core, ai)
    print(OUT_PATH)
    print(core)
    print(ai)


if __name__ == "__main__":
    main()
